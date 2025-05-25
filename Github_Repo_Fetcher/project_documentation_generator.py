#!/usr/bin/env python3

import os
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import re
from datetime import datetime
import glob
import openai
from dataclasses import dataclass
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
import git
from file_classifier import FileClassifier
import time

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@dataclass
class FileAnalysis:
    """Data class to store file analysis information."""
    path: str
    content: str
    language: str
    framework: Optional[str] = None
    purpose: Optional[str] = None
    dependencies: List[str] = None
    complexity: float = 0.0

class ProjectDocumentationGenerator:
    """Generate comprehensive project documentation using ChatGPT-4."""
    
    def __init__(self, repo_path: str):
        """
        Initialize the documentation generator.
        
        Args:
            repo_path (str): Path to the repository to analyze
        """
        self.repo_path = Path(repo_path)
        if not self.repo_path.exists():
            raise ValueError(f"Repository path does not exist: {repo_path}")
        
        # Create documentation directory
        self.docs_dir = self.repo_path / 'Logic Understanding'
        self.docs_dir.mkdir(exist_ok=True)
        
        # Initialize OpenAI client
        self._setup_openai()
        
        # Initialize Git repository
        try:
            self.repo = git.Repo(self.repo_path)
        except git.InvalidGitRepositoryError:
            logger.warning("Not a git repository. Versioning will be limited.")
            self.repo = None
        
        # Initialize file classifier
        self.classifier = FileClassifier(str(self.repo_path))
        
        # Store analysis results
        self.file_analyses: List[FileAnalysis] = []
        self.project_summary: Dict = {}
        self.architecture_overview: Dict = {}
        
    def _setup_openai(self):
        """Setup OpenAI API with proper error handling."""
        # Get the directory containing the script
        script_dir = Path(__file__).parent.absolute()
        env_path = script_dir / '.env'
        
        if not env_path.exists():
            raise FileNotFoundError(f".env file not found at: {env_path}")
        
        load_dotenv(env_path)
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise ValueError("OPENAI_API_KEY not found in .env file")
        
        if not api_key.startswith('sk-'):
            logger.warning("OPENAI_API_KEY doesn't appear to be in the correct format")
        
        self.client = openai.OpenAI(api_key=api_key)
        logger.info("Successfully initialized OpenAI client")

    def _analyze_file(self, file_path: Path) -> Optional[FileAnalysis]:
        """Analyze a single file and return its analysis."""
        try:
            # Skip binary files and large files
            if not file_path.is_file() or self._is_binary_file(file_path):
                return None
            
            # Get file classification
            classification = self.classifier.classify_file(str(file_path))
            
            # Read file content
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            
            # Skip if file is too large
            if len(content) > 100000:  # Skip files larger than 100KB
                logger.warning(f"File too large to analyze: {file_path}")
                return None
            
            # Analyze file with ChatGPT
            analysis = self._get_file_analysis(content, classification)
            
            return FileAnalysis(
                path=str(file_path.relative_to(self.repo_path)),
                content=content,
                language=classification.get('language', 'unknown'),
                framework=classification.get('framework'),
                purpose=analysis.get('purpose'),
                dependencies=analysis.get('dependencies', []),
                complexity=analysis.get('complexity', 0.0)
            )
            
        except Exception as e:
            logger.error(f"Error analyzing file {file_path}: {str(e)}")
            return None

    def _is_binary_file(self, file_path: Path) -> bool:
        """Check if a file is binary."""
        try:
            with open(file_path, 'rb') as f:
                chunk = f.read(1024)
                return b'\0' in chunk
        except Exception:
            return True

    def _get_file_analysis(self, content: str, classification: Dict) -> Dict:
        """Get file analysis from ChatGPT."""
        try:
            prompt = f"""Analyze this {classification.get('language', 'code')} file and provide a JSON object with these exact keys:
{{
    "purpose": "string describing the main purpose of this file",
    "dependencies": ["list", "of", "dependencies"],
    "complexity": float between 0.0 and 1.0
}}

File content:
{content[:4000]}  # Limit content to avoid token limits

IMPORTANT: Return ONLY a valid JSON object with the exact keys shown above. Do not include any other text or explanation."""

            response = self.client.chat.completions.create(
                model="gpt-4-0125-preview",
                messages=[
                    {"role": "system", "content": "You are an expert code analyzer. Return ONLY valid JSON objects with the exact keys specified. Do not include any other text or explanation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=1000,
                response_format={ "type": "json_object" }
            )
            
            try:
                analysis = json.loads(response.choices[0].message.content)
                # Validate required keys
                required_keys = {'purpose', 'dependencies', 'complexity'}
                if not all(key in analysis for key in required_keys):
                    raise ValueError("Missing required keys in response")
                return analysis
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON in file analysis response: {str(e)}")
                logger.error(f"Raw response: {response.choices[0].message.content}")
                return {
                    'purpose': 'Analysis failed - invalid JSON response',
                    'dependencies': [],
                    'complexity': 0.0
                }
            
        except Exception as e:
            logger.error(f"Error getting file analysis: {str(e)}")
            return {
                'purpose': 'Analysis failed',
                'dependencies': [],
                'complexity': 0.0
            }

    def _find_app_directory(self) -> Optional[Path]:
        """Find the main application directory."""
        possible_app_dirs = ['app', 'src', 'application', 'apps']
        for dir_name in possible_app_dirs:
            app_dir = self.repo_path / dir_name
            if app_dir.exists() and app_dir.is_dir():
                return app_dir
        return None

    def _estimate_tokens(self, text: str) -> int:
        """Estimate the number of tokens in a text string."""
        # Rough estimation: 1 token ≈ 4 characters for English text
        return len(text) // 4

    def _analyze_project_structure(self, app_dir: Path) -> Dict:
        """Analyze the project structure to understand its organization."""
        try:
            structure = {
                'framework': None,
                'architecture': None,
                'main_components': [],
                'entry_points': [],
                'config_files': []
            }
            
            # Detect framework and architecture
            if (app_dir / 'app/Http/Controllers').exists():
                structure['framework'] = 'Laravel'
                structure['architecture'] = 'MVC'
                structure['main_components'] = [
                    'Controllers (app/Http/Controllers)',
                    'Models (app/Models)',
                    'Views (resources/views)',
                    'Routes (routes)',
                    'Migrations (database/migrations)'
                ]
            elif (app_dir / 'app/controllers').exists():
                structure['framework'] = 'Ruby on Rails'
                structure['architecture'] = 'MVC'
            elif (app_dir / 'src/main/java').exists():
                structure['framework'] = 'Spring Boot'
                structure['architecture'] = 'Layered Architecture'
            elif (app_dir / 'src/components').exists():
                structure['framework'] = 'React'
                structure['architecture'] = 'Component-based'
            
            # Find entry points
            for pattern in ['*.php', '*.py', '*.js', '*.java', 'app.py', 'index.js', 'main.py']:
                for file in app_dir.glob(f'**/{pattern}'):
                    if any(x in str(file) for x in ['public', 'bin', 'scripts', 'entry']):
                        structure['entry_points'].append(str(file.relative_to(app_dir)))
            
            # Find config files
            for pattern in ['*.config.*', '*.env*', '*.yml', '*.yaml', '*.json', '*.xml']:
                for file in app_dir.glob(f'**/{pattern}'):
                    if 'config' in str(file).lower() or 'settings' in str(file).lower():
                        structure['config_files'].append(str(file.relative_to(app_dir)))
            
            return structure
            
        except Exception as e:
            logger.error(f"Error analyzing project structure: {str(e)}")
            return {
                'framework': None,
                'architecture': None,
                'main_components': [],
                'entry_points': [],
                'config_files': []
            }

    def _process_batch(self, files_batch: List[Dict], project_structure: Dict) -> List[Dict]:
        """Process a batch of files and get their analysis."""
        try:
            # Prepare the files content for analysis
            files_info = "\n\n".join([
                f"File: {file['path']}\n"
                f"Language: {file['language']}\n"
                f"Framework: {file['framework'] or 'None'}\n"
                f"Content:\n{file['content']}\n"
                f"{'='*80}\n"
                for file in files_batch
            ])
            
            # Add project structure context
            structure_info = f"""
Project Structure:
Framework: {project_structure['framework'] or 'Unknown'}
Architecture: {project_structure['architecture'] or 'Unknown'}
Main Components: {', '.join(project_structure['main_components'])}
Entry Points: {', '.join(project_structure['entry_points'])}
Config Files: {', '.join(project_structure['config_files'])}
"""
            
            prompt = f"""Analyze these code files and provide a JSON object with these exact keys:
{{
    "project_description": "string describing the project's purpose, goals, target users, and key problems solved",
    "core_functionality": ["list", "of", "strings describing core functional components"],
    "key_features": ["list", "of", "strings describing key features"]
}}

Project Structure:
{structure_info}

Files to analyze:
{files_info}

IMPORTANT: Return ONLY a valid JSON object with the exact keys shown above. Do not include any other text or explanation."""

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert code analyzer and technical writer. Return ONLY valid JSON objects with the exact keys specified. Do not include any other text or explanation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2000,
                response_format={ "type": "json_object" }
            )
            
            # Add a small delay to respect rate limits
            time.sleep(1)
            
            try:
                # Parse the response and validate its structure
                analysis = json.loads(response.choices[0].message.content)
                
                # Ensure the response has the required structure
                if not isinstance(analysis, dict):
                    raise ValueError("Response is not a dictionary")
                
                if 'project_description' not in analysis or not isinstance(analysis['project_description'], str):
                    analysis['project_description'] = "Failed to generate project description"
                
                if 'core_functionality' not in analysis or not isinstance(analysis['core_functionality'], list):
                    analysis['core_functionality'] = []
                else:
                    # Ensure all core_functionality items are strings
                    analysis['core_functionality'] = [
                        str(func) if not isinstance(func, str) else func
                        for func in analysis['core_functionality']
                    ]
                
                if 'key_features' not in analysis or not isinstance(analysis['key_features'], list):
                    analysis['key_features'] = []
                else:
                    # Ensure all key_features are strings
                    analysis['key_features'] = [
                        str(feature) if not isinstance(feature, str) else feature
                        for feature in analysis['key_features']
                    ]
                
                return [analysis]
                
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON response: {str(e)}")
                return [{
                    'project_description': 'Failed to parse analysis response',
                    'core_functionality': [],
                    'key_features': []
                }]
            
        except Exception as e:
            logger.error(f"Error processing batch: {str(e)}")
            return [{
                'project_description': f'Analysis failed for batch: {str(e)}',
                'core_functionality': [],
                'key_features': []
            }]

    def _get_file_classification(self, file_path: Path) -> Dict:
        """Get file classification with proper error handling."""
        try:
            # Basic classification based on file extension
            ext = file_path.suffix.lower()
            language_map = {
                '.php': 'PHP',
                '.py': 'Python',
                '.js': 'JavaScript',
                '.ts': 'TypeScript',
                '.java': 'Java',
                '.cpp': 'C++',
                '.c': 'C',
                '.h': 'C/C++ Header',
                '.hpp': 'C++ Header',
                '.cs': 'C#',
                '.go': 'Go',
                '.rb': 'Ruby',
                '.jsx': 'React',
                '.tsx': 'React TypeScript'
            }
            
            # Framework detection based on file path and content
            framework = None
            if ext == '.php':
                if 'Laravel' in str(file_path) or 'app/Http/Controllers' in str(file_path):
                    framework = 'Laravel'
                elif 'Symfony' in str(file_path):
                    framework = 'Symfony'
            elif ext in ['.js', '.jsx', '.ts', '.tsx']:
                if 'react' in str(file_path).lower() or 'components' in str(file_path).lower():
                    framework = 'React'
                elif 'angular' in str(file_path).lower():
                    framework = 'Angular'
                elif 'vue' in str(file_path).lower():
                    framework = 'Vue.js'
            elif ext == '.py':
                if 'django' in str(file_path).lower():
                    framework = 'Django'
                elif 'flask' in str(file_path).lower():
                    framework = 'Flask'
            
            return {
                'language': language_map.get(ext, 'Unknown'),
                'framework': framework
            }
            
        except Exception as e:
            logger.warning(f"Error classifying file {file_path}: {str(e)}")
            return {
                'language': 'Unknown',
                'framework': None
            }

    def analyze_repository(self):
        """Analyze the application code in the app directory."""
        logger.info(f"Starting repository analysis: {self.repo_path}")
        
        # Find app directory
        app_dir = self._find_app_directory()
        if not app_dir:
            raise ValueError("Could not find application directory (app, src, or application folder)")
        
        logger.info(f"Found application directory: {app_dir}")
        
        # Analyze project structure
        project_structure = self._analyze_project_structure(app_dir)
        logger.info(f"Detected framework: {project_structure['framework']}")
        
        # Get list of code files to analyze
        files_to_analyze = []
        for ext in ['.py', '.js', '.java', '.cpp', '.c', '.h', '.hpp', '.cs', '.go', '.rb', '.php', '.ts', '.jsx', '.tsx']:
            files_to_analyze.extend(app_dir.glob(f"**/*{ext}"))
        
        # Collect file contents and basic analysis
        files_content = []
        for file_path in files_to_analyze:
            try:
                if not file_path.is_file() or self._is_binary_file(file_path):
                    continue
                
                # Get file classification using our new method
                classification = self._get_file_classification(file_path)
                
                # Read file content
                try:
                    content = file_path.read_text(encoding='utf-8', errors='ignore')
                except Exception as e:
                    logger.warning(f"Error reading file {file_path}: {str(e)}")
                    continue
                
                # Skip if file is too large
                if len(content) > 100000:  # Skip files larger than 100KB
                    logger.warning(f"File too large to analyze: {file_path}")
                    continue
                
                files_content.append({
                    'path': str(file_path.relative_to(self.repo_path)),
                    'content': content,
                    'language': classification['language'],
                    'framework': classification['framework']
                })
                
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {str(e)}")
                continue
        
        if not files_content:
            raise ValueError("No code files found to analyze in the application directory")
        
        # Get detailed code logic analysis
        self.code_analysis = self._get_code_logic_analysis(files_content, project_structure)
        
        logger.info(f"Repository analysis complete. Analyzed {len(files_content)} files")

    def _create_word_document(self) -> Document:
        """Create a comprehensive Word document from the code analysis."""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Project Documentation"
        doc.core_properties.author = "Documentation Generator"
        
        # Add title
        title = doc.add_heading('Project Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add version information
        if self.repo:
            version = doc.add_paragraph()
            version.add_run('Version: ').bold = True
            version.add_run(f"{self.repo.head.commit.hexsha[:8]}\n")
            version.add_run('Last Updated: ').bold = True
            version.add_run(f"{datetime.fromtimestamp(self.repo.head.commit.committed_date).strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add project description
        doc.add_heading('Project Description', level=1)
        description = doc.add_paragraph()
        description.add_run(self.code_analysis['project_description'])
        
        # Add core functionality
        doc.add_heading('Core Functionality', level=1)
        core_func = doc.add_paragraph()
        for func in self.code_analysis['core_functionality']:
            core_func.add_run(f"• {func}\n")
        
        # Add key features
        doc.add_heading('Key Features', level=1)
        features = doc.add_paragraph()
        for feature in self.code_analysis['key_features']:
            features.add_run(f"• {feature}\n")
        
        return doc

    def generate_documentation(self) -> str:
        """Generate and save the documentation."""
        logger.info("Generating documentation...")
        
        # Analyze repository
        self.analyze_repository()
        
        # Create Word document
        doc = self._create_word_document()
        
        # Generate filename with version
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if self.repo:
            version = self.repo.head.commit.hexsha[:8]
            filename = f"project_logic_documentation_v1_{version}_{timestamp}.docx"
        else:
            filename = f"project_logic_documentation_v1_{timestamp}.docx"
        
        # Save document
        doc_path = self.docs_dir / filename
        doc.save(doc_path)
        
        # Create latest symlink/copy
        latest_doc = self.docs_dir / "project_logic_documentation_latest.docx"
        if os.name == 'nt':
            shutil.copy2(doc_path, latest_doc)
        else:
            if latest_doc.exists():
                latest_doc.unlink()
            latest_doc.symlink_to(doc_path)
        
        logger.info(f"Documentation saved to {doc_path}")
        return filename

    def _get_code_logic_analysis(self, files_content: List[Dict], project_structure: Dict) -> Dict:
        """Get detailed code logic analysis from ChatGPT, processing files in batches."""
        try:
            # Sort files by size to optimize batching
            files_content.sort(key=lambda x: len(x['content']))
            
            # Process files in batches to stay within token limits
            batch_size = 5  # Start with a small batch size
            max_tokens_per_request = 8000  # Leave room for response
            current_batch = []
            current_batch_tokens = 0
            all_analyses = []
            
            for file in files_content:
                file_tokens = self._estimate_tokens(file['content'])
                
                # If a single file is too large, process it alone
                if file_tokens > max_tokens_per_request:
                    logger.warning(f"File {file['path']} is too large, processing separately")
                    if current_batch:
                        all_analyses.extend(self._process_batch(current_batch, project_structure))
                        current_batch = []
                        current_batch_tokens = 0
                    
                    # Process large file with truncated content
                    truncated_file = file.copy()
                    truncated_file['content'] = file['content'][:30000]  # Truncate to ~7500 tokens
                    all_analyses.extend(self._process_batch([truncated_file], project_structure))
                    continue
                
                # Add file to current batch if it fits
                if current_batch_tokens + file_tokens < max_tokens_per_request and len(current_batch) < batch_size:
                    current_batch.append(file)
                    current_batch_tokens += file_tokens
                else:
                    # Process current batch and start a new one
                    if current_batch:
                        all_analyses.extend(self._process_batch(current_batch, project_structure))
                    current_batch = [file]
                    current_batch_tokens = file_tokens
            
            # Process any remaining files
            if current_batch:
                all_analyses.extend(self._process_batch(current_batch, project_structure))
            
            # Combine all analyses
            return self._combine_analyses(all_analyses, project_structure)
            
        except Exception as e:
            logger.error(f"Error getting code logic analysis: {str(e)}")
            return {
                'project_description': 'Analysis failed',
                'core_functionality': [],
                'key_features': []
            }

    def _combine_analyses(self, analyses: List[Dict], project_structure: Dict) -> Dict:
        """Combine multiple analyses into a single coherent analysis."""
        try:
            # Prepare combined analysis data
            combined_data = {
                'project_description': [],
                'core_functionality': [],  # Added core functionality
                'key_features': []
            }
            
            # Combine all analyses
            for analysis in analyses:
                # Handle project description
                if isinstance(analysis.get('project_description'), str) and analysis['project_description'] != 'Analysis failed':
                    combined_data['project_description'].append(analysis['project_description'])
                
                # Handle core functionality
                core_funcs = analysis.get('core_functionality', [])
                if isinstance(core_funcs, list):
                    combined_data['core_functionality'].extend([
                        str(func) for func in core_funcs
                        if func and str(func).strip()
                    ])
                elif isinstance(core_funcs, (dict, str)):
                    combined_data['core_functionality'].append(str(core_funcs))
                
                # Handle key features
                features = analysis.get('key_features', [])
                if isinstance(features, list):
                    combined_data['key_features'].extend([
                        str(feature) for feature in features
                        if feature and str(feature).strip()
                    ])
                elif isinstance(features, (dict, str)):
                    combined_data['key_features'].append(str(features))
            
            # Remove duplicates while preserving order
            for key in ['core_functionality', 'key_features']:
                seen = set()
                combined_data[key] = [
                    x for x in combined_data[key]
                    if not (x in seen or seen.add(x))
                ]
            
            # If no valid descriptions were found, add a default message
            if not combined_data['project_description']:
                combined_data['project_description'].append(
                    "Unable to generate project description from the analyzed files."
                )
            
            # Get final combined analysis from GPT-4
            structure_info = f"""
Project Structure:
Framework: {project_structure['framework'] or 'Unknown'}
Architecture: {project_structure['architecture'] or 'Unknown'}
Main Components: {', '.join(project_structure['main_components'])}
Entry Points: {', '.join(project_structure['entry_points'])}
Config Files: {', '.join(project_structure['config_files'])}
"""
            
            prompt = f"""Based on these partial analyses and project structure, provide a JSON object with these exact keys:
{{
    "project_description": "string with comprehensive explanation of the project's purpose, goals, target users, and key problems solved",
    "core_functionality": ["list", "of", "strings describing core functional components"],
    "key_features": ["list", "of", "strings describing key features"]
}}

Project Structure:
{structure_info}

Partial analyses:
{json.dumps(combined_data, indent=2)}

IMPORTANT: Return ONLY a valid JSON object with the exact keys shown above. Do not include any other text or explanation."""

            response = self.client.chat.completions.create(
                model="gpt-4-0125-preview",
                messages=[
                    {"role": "system", "content": "You are an expert technical writer. Return ONLY valid JSON objects with the exact keys specified. Do not include any other text or explanation."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2000,
                response_format={ "type": "json_object" }
            )
            
            try:
                final_analysis = json.loads(response.choices[0].message.content)
                
                # Validate the final analysis structure
                if not isinstance(final_analysis, dict):
                    raise ValueError("Final analysis is not a dictionary")
                
                if 'project_description' not in final_analysis or not isinstance(final_analysis['project_description'], str):
                    final_analysis['project_description'] = "Failed to generate final project description"
                
                if 'core_functionality' not in final_analysis or not isinstance(final_analysis['core_functionality'], list):
                    final_analysis['core_functionality'] = []
                else:
                    final_analysis['core_functionality'] = [
                        str(func) if not isinstance(func, str) else func
                        for func in final_analysis['core_functionality']
                    ]
                
                if 'key_features' not in final_analysis or not isinstance(final_analysis['key_features'], list):
                    final_analysis['key_features'] = []
                else:
                    final_analysis['key_features'] = [
                        str(feature) if not isinstance(feature, str) else feature
                        for feature in final_analysis['key_features']
                    ]
                
                return final_analysis
                
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON in final analysis response: {str(e)}")
                logger.error(f"Raw response: {response.choices[0].message.content}")
                return {
                    'project_description': 'Analysis failed - invalid JSON response',
                    'core_functionality': combined_data.get('core_functionality', []),
                    'key_features': combined_data.get('key_features', [])
                }
            
        except Exception as e:
            logger.error(f"Error combining analyses: {str(e)}")
            return {
                'project_description': 'Analysis failed during combination',
                'core_functionality': combined_data.get('core_functionality', []),
                'key_features': combined_data.get('key_features', [])
            }

def main():
    """Main function to generate project documentation."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate project logic documentation')
    parser.add_argument('repo_path', help='Path to the repository to analyze')
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    
    try:
        generator = ProjectDocumentationGenerator(args.repo_path)
        doc_file = generator.generate_documentation()
        
        print("\nProject Logic Documentation Generated Successfully!")
        print("=" * 50)
        print(f"Documentation file: Logic Understanding/{doc_file}")
        print(f"Latest version: Logic Understanding/project_logic_documentation_latest.docx")
        print("\nThe documentation includes:")
        print("• Project Description")
        print("• Core Functionality")
        print("• Key Features")
        
    except Exception as e:
        logger.error(f"Error generating documentation: {str(e)}")
        raise

if __name__ == '__main__':
    main() 