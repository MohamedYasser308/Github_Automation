#!/usr/bin/env python3

import os
import re
import json
import logging
import argparse
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import git

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class RouteParser:
    """Parse Laravel route files directly."""
    
    def __init__(self, repo_path: str):
        self.repo_path = Path(repo_path).resolve()
        if not self.repo_path.exists():
            raise ValueError(f"Repository path does not exist: {repo_path}")
        
        # Create documentation directory structure
        self.docs_dir = self.repo_path / 'API Documentation'
        self.versions_dir = self.docs_dir / 'versions'
        self.latest_dir = self.docs_dir / 'latest'
        
        # Create directories if they don't exist
        self.docs_dir.mkdir(exist_ok=True)
        self.versions_dir.mkdir(exist_ok=True)
        self.latest_dir.mkdir(exist_ok=True)
        
        # Initialize Git repository for versioning
        try:
            self.repo = git.Repo(self.repo_path)
            self.current_branch = self.repo.active_branch.name
        except git.InvalidGitRepositoryError:
            logger.warning("Not a git repository. Versioning will be limited.")
            self.repo = None
            self.current_branch = 'unknown'
    
    def _find_route_files(self) -> List[Path]:
        """Find all route files in the repository."""
        route_files = []
        routes_dir = self.repo_path / 'routes'
        
        if not routes_dir.exists():
            logger.error(f"routes directory not found at {routes_dir}")
            return []
        
        logger.info(f"Searching for route files in {routes_dir}")
        php_files = list(routes_dir.glob('*.php'))
        
        if php_files:
            logger.info(f"Found {len(php_files)} PHP route files:")
            for file in php_files:
                logger.info(f"  - {file.relative_to(self.repo_path)}")
            route_files.extend(php_files)
        
        return route_files

    def _parse_route_file(self, file_path: Path) -> List[Dict]:
        """Parse a single route file and extract route information."""
        routes = []
        content = self._read_file_content(file_path)
        
        if not content:
            logger.warning(f"Empty file: {file_path}")
            return []
            
        logger.debug(f"Parsing routes in {file_path}")
        
        # Common route patterns with Laravel 8+ syntax
        patterns = [
            # Route::get/post/put/delete pattern with array syntax
            r'Route::(get|post|put|delete|patch|options)\s*\(\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*,\s*[\'"]([^\'"]+)[\'"]\s*\]\s*\)',
            # Route::match pattern with array syntax
            r'Route::match\s*\(\s*\[([^\]]+)\]\s*,\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*,\s*[\'"]([^\'"]+)[\'"]\s*\]\s*\)',
            # Route::any pattern with array syntax
            r'Route::any\s*\(\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*,\s*[\'"]([^\'"]+)[\'"]\s*\]\s*\)',
            # Route::resource pattern with array syntax
            r'Route::resource\s*\(\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*\]\s*\)',
            # Route group pattern
            r'Route::group\s*\(\s*\[([^\]]+)\]\s*,\s*function\s*\(\s*\)\s*{([^}]+)}\)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.MULTILINE | re.DOTALL)
            for match in matches:
                logger.debug(f"Found route match: {match.group(0)}")
                route_info = self._extract_route_info(match, pattern)
                if route_info:
                    if isinstance(route_info, list):
                        routes.extend(route_info)
                    else:
                        routes.append(route_info)
        
        if not routes:
            logger.warning(f"No routes found in {file_path}")
        else:
            logger.info(f"Found {len(routes)} routes in {file_path}")
            for route in routes:
                logger.debug(f"Route: {route['methods']} {route['path']} -> {route['controller']}")
        
        return routes

    def _extract_route_info(self, match: re.Match, pattern: str) -> Optional[Dict]:
        """Extract route information from a regex match."""
        try:
            if 'Route::group' in pattern:
                # Handle route groups
                middleware_str = match.group(1)
                group_content = match.group(2)
                
                # Extract middleware from group
                middleware = []
                middleware_match = re.search(r'middleware\s*\(\s*\[([^\]]+)\]', middleware_str)
                if middleware_match:
                    middleware = [m.strip().strip("'\"") for m in middleware_match.group(1).split(',')]
                
                # Parse routes inside the group
                group_routes = []
                for route_pattern in [
                    r'Route::(get|post|put|delete|patch|options)\s*\(\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*,\s*[\'"]([^\'"]+)[\'"]\s*\]\s*\)',
                    r'Route::match\s*\(\s*\[([^\]]+)\]\s*,\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*,\s*[\'"]([^\'"]+)[\'"]\s*\]\s*\)',
                    r'Route::any\s*\(\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*,\s*[\'"]([^\'"]+)[\'"]\s*\]\s*\)',
                    r'Route::resource\s*\(\s*[\'"]([^\'"]+)[\'"]\s*,\s*\[([^,]+)::class\s*\]\s*\)'
                ]:
                    for route_match in re.finditer(route_pattern, group_content, re.MULTILINE | re.DOTALL):
                        route_info = self._extract_route_info(route_match, route_pattern)
                        if route_info:
                            if isinstance(route_info, list):
                                for r in route_info:
                                    r['middleware'].extend(middleware)
                                group_routes.extend(route_info)
                            else:
                                route_info['middleware'].extend(middleware)
                                group_routes.append(route_info)
                
                return group_routes if group_routes else None
            
            if 'Route::match' in pattern:
                methods = [m.strip().upper() for m in match.group(1).split(',')]
                path = match.group(2)
                controller = match.group(3).strip()
                method = match.group(4).strip("'\"")
            elif 'Route::any' in pattern:
                methods = ['GET', 'POST', 'PUT', 'PATCH', 'DELETE', 'OPTIONS']
                path = match.group(1)
                controller = match.group(2).strip()
                method = match.group(3).strip("'\"")
            elif 'Route::resource' in pattern:
                base_path = match.group(1)
                controller = match.group(2).strip()
                return self._generate_resource_routes(base_path, controller)
            else:
                methods = [match.group(1).upper()]
                path = match.group(2)
                controller = match.group(3).strip()
                method = match.group(4).strip("'\"")
            
            # Clean up the controller name
            controller = controller.replace('::class', '')
            
            # Extract middleware if present
            middleware = []
            middleware_match = re.search(r'middleware\s*\(\s*\[([^\]]+)\]', path)
            if middleware_match:
                middleware = [m.strip().strip("'\"") for m in middleware_match.group(1).split(',')]
            
            # Extract route name if present
            name = None
            name_match = re.search(r'name\s*\(\s*[\'"]([^\'"]+)[\'"]', path)
            if name_match:
                name = name_match.group(1)
            
            route_info = {
                "methods": methods,
                "path": path,
                "controller": f"{controller}@{method}",
                "middleware": middleware,
                "name": name,
                "description": self._generate_route_description(methods, path, controller, method),
                "parameters": self._extract_route_parameters(path),
                "auth_required": any('auth' in m.lower() for m in middleware),
                "rate_limit": self._extract_rate_limit(middleware)
            }
            
            logger.debug(f"Extracted route info: {route_info}")
            return [route_info]
            
        except Exception as e:
            logger.warning(f"Error extracting route info: {str(e)}")
            logger.debug(f"Match groups: {match.groups()}")
            return None

    def _generate_resource_routes(self, base_path: str, controller: str) -> List[Dict]:
        """Generate standard REST resource routes."""
        resource_routes = []
        standard_methods = {
            'index': ['GET'],
            'store': ['POST'],
            'show': ['GET'],
            'update': ['PUT', 'PATCH'],
            'destroy': ['DELETE']
        }
        
        for method_name, http_methods in standard_methods.items():
            path = base_path if method_name == 'index' else f"{base_path}/{{id}}"
            resource_routes.append({
                "methods": http_methods,
                "path": path,
                "controller": f"{controller}@{method_name}",
                "middleware": [],
                "name": f"{base_path}.{method_name}",
                "description": self._generate_route_description(http_methods, path, controller, method_name),
                "parameters": self._extract_route_parameters(path),
                "auth_required": False,
                "rate_limit": None
            })
        
        return resource_routes

    def _extract_route_parameters(self, path: str) -> List[Dict]:
        """Extract parameters from route path."""
        parameters = []
        param_pattern = r'{([^}]+)}'
        
        for param in re.finditer(param_pattern, path):
            param_name = param.group(1)
            # Remove any parameter constraints
            param_name = param_name.split(':')[0]
            parameters.append({
                "name": param_name,
                "type": "string",  # Default type
                "required": True,
                "description": f"Parameter for {param_name}"
            })
        
        return parameters

    def _extract_rate_limit(self, middleware: List[str]) -> Optional[str]:
        """Extract rate limit information from middleware."""
        for m in middleware:
            if 'throttle' in m.lower():
                return m
        return None

    def _generate_route_description(self, methods: List[str], path: str, controller: str, method: str) -> str:
        """Generate a description for the route based on its components."""
        method_str = '/'.join(methods)
        return f"{method_str} {path} - Handled by {controller}"

    def _read_file_content(self, file_path: Path) -> str:
        """Read and return the content of a file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                logger.debug(f"Read {len(content)} bytes from {file_path}")
                return content
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            return ""

    def _create_api_document(self, documentation: Dict) -> Document:
        """Create the API documentation document."""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "API Documentation"
        doc.core_properties.author = "API Documentation Generator"
        
        # Add title
        title = doc.add_heading('API Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add version information
        if self.repo:
            version = doc.add_paragraph()
            version.add_run('Version: ').bold = True
            version.add_run(f"{self.repo.head.commit.hexsha[:8]}\n")
            version.add_run('Last Updated: ').bold = True
            version.add_run(f"{datetime.fromtimestamp(self.repo.head.commit.committed_date).strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add API Overview
        doc.add_heading('API Overview', level=1)
        overview = doc.add_paragraph(documentation['api_overview'])
        
        # Add endpoints by file
        for endpoint_group in documentation['endpoints']:
            doc.add_heading(f'Endpoints in {endpoint_group["file"]}', level=1)
            
            for route in endpoint_group['routes']:
                # Endpoint Header
                methods = '/'.join(route['methods'])
                header = doc.add_heading(f"{methods} {route['path']}", level=2)
                header.style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                
                # Controller
                p = doc.add_paragraph()
                p.add_run('Controller: ').bold = True
                p.add_run(route['controller'])
                
                # Description
                p = doc.add_paragraph()
                p.add_run('Description: ').bold = True
                p.add_run(route['description'])
                
                # Parameters
                if route['parameters']:
                    doc.add_heading('Parameters', level=3)
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    # Add header row
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Name'
                    header_cells[1].text = 'Type'
                    header_cells[2].text = 'Required'
                    header_cells[3].text = 'Description'
                    
                    # Add parameter rows
                    for param in route['parameters']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = param['name']
                        row_cells[1].text = param['type']
                        row_cells[2].text = 'Yes' if param['required'] else 'No'
                        row_cells[3].text = param['description']
                
                # Authentication
                if route.get('auth_required'):
                    p = doc.add_paragraph()
                    p.add_run('Authentication Required: ').bold = True
                    p.add_run('Yes')
                
                # Rate Limiting
                if route.get('rate_limit'):
                    p = doc.add_paragraph()
                    p.add_run('Rate Limit: ').bold = True
                    p.add_run(str(route['rate_limit']))
                
                # Example Usage
                example = self._generate_example_usage(route)
                if example:
                    doc.add_heading('Example Usage', level=3)
                    # Create a paragraph with monospace font for code
                    p = doc.add_paragraph()
                    run = p.add_run(example)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    # Add a light gray background
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.left_indent = Pt(12)
                    p.paragraph_format.right_indent = Pt(12)
                
                doc.add_paragraph()  # Add spacing between endpoints
        
        return doc

    def _generate_example_usage(self, route: Dict) -> str:
        """Generate example usage for a route."""
        method = route['methods'][0]  # Use first method for example
        path = route['path']
        params = {p['name']: f"<{p['name']}>" for p in route['parameters']}
        
        # Replace path parameters
        for param_name, value in params.items():
            path = path.replace(f"{{{param_name}}}", value)
        
        # Generate curl command
        curl_cmd = f"curl -X {method} \\\n"
        if route.get('auth_required'):
            curl_cmd += "  -H 'Authorization: Bearer <your_token>' \\\n"
        curl_cmd += f"  'http://your-domain.com{path}'"
        
        return curl_cmd

    def generate_api_documentation(self) -> str:
        """Generate and save the API documentation."""
        logger.info("Generating API documentation...")
        
        # Find route files
        route_files = self._find_route_files()
        if not route_files:
            raise ValueError("No route files found in the repository")
        
        # Parse all route files
        all_routes = []
        for file_path in route_files:
            routes = self._parse_route_file(file_path)
            if routes:
                # Flatten the list of routes (in case some are nested)
                flat_routes = []
                for route in routes:
                    if isinstance(route, list):
                        flat_routes.extend(route)
                    else:
                        flat_routes.append(route)
                
                if flat_routes:  # Only add if we have routes
                    all_routes.append({
                        "file": str(file_path.relative_to(self.repo_path)),
                        "routes": flat_routes
                    })
        
        if not all_routes:
            raise ValueError("No valid routes found in any route files")
        
        # Create documentation structure
        documentation = {
            "api_overview": "API Documentation generated from Laravel route files",
            "endpoints": all_routes
        }
        
        # Create API document
        doc = self._create_api_document(documentation)
        
        # Get version information
        version_info = self._get_version_info()
        
        # Generate filenames
        version_filename = f"api_documentation_v{version_info['version']}_{version_info['branch']}_{version_info['commit']}_{version_info['timestamp']}.docx"
        latest_filename = "api_documentation_latest.docx"
        
        # Save versioned document
        version_path = self.versions_dir / version_filename
        doc.save(version_path)
        
        # Save latest version
        latest_path = self.latest_dir / latest_filename
        if os.name == 'nt':
            shutil.copy2(version_path, latest_path)
        else:
            if latest_path.exists():
                latest_path.unlink()
            latest_path.symlink_to(version_path)
        
        # Save analysis as JSON
        analysis_path = self.versions_dir / f"{version_filename}.json"
        with open(analysis_path, 'w') as f:
            json.dump({
                'version_info': version_info,
                'documentation': documentation
            }, f, indent=2)
        
        logger.info(f"API documentation saved to {version_path}")
        logger.info(f"Latest version available at {latest_path}")
        logger.info(f"Analysis data saved to {analysis_path}")
        
        return version_filename

    def _get_version_info(self) -> Dict[str, str]:
        """Get version information for the documentation."""
        version_info = {
            'timestamp': datetime.now().strftime("%Y%m%d_%H%M%S"),
            'branch': self.current_branch,
            'commit': 'unknown',
            'version': '1.0.0'  # Default version
        }
        
        if self.repo:
            try:
                commit = self.repo.head.commit
                version_info.update({
                    'commit': commit.hexsha[:8],
                    'commit_date': datetime.fromtimestamp(commit.committed_date).strftime('%Y-%m-%d %H:%M:%S'),
                    'commit_message': commit.message.split('\n')[0],
                    'author': f"{commit.author.name} <{commit.author.email}>"
                })
                
                # Try to get version from git tags
                tags = [tag for tag in self.repo.tags if tag.commit == commit]
                if tags:
                    version_info['version'] = tags[0].name
            except Exception as e:
                logger.warning(f"Error getting git version info: {str(e)}")
        
        return version_info

def main():
    """Main function to generate API documentation."""
    parser = argparse.ArgumentParser(
        description='Generate API documentation from Laravel route files',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Generate API documentation
  python api_documentation_generator.py /path/to/repo
  
  # Generate with verbose logging
  python api_documentation_generator.py /path/to/repo -v
        """
    )
    
    parser.add_argument('repo_path', help='Path to the repository')
    parser.add_argument('--verbose', '-v', action='store_true',
                      help='Enable verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    
    try:
        parser = RouteParser(args.repo_path)
        doc_file = parser.generate_api_documentation()
        
        print("\nAPI Documentation Generated Successfully!")
        print("=" * 50)
        print(f"Documentation file: API Documentation/versions/{doc_file}")
        print(f"Latest version: API Documentation/latest/api_documentation_latest.docx")
        print("\nThe documentation includes:")
        print("• API Overview")
        print("• Endpoints by file")
        print("• Detailed endpoint documentation")
        print("• Parameters and return values")
        print("• Usage examples")
        print("\nAnalysis data is stored in:")
        print(f"• API Documentation/versions/{doc_file}.json")
        
    except ValueError as e:
        logger.error(f"Configuration error: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"Error generating API documentation: {str(e)}")
        raise

if __name__ == '__main__':
    main() 