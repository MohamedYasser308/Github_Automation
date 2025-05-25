#!/usr/bin/env python3

import os
import json
import logging
import shutil
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import openai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from dotenv import load_dotenv
import git
import time

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class UATDocumentationGenerator:
    """Generate UAT documentation from project logic documentation."""
    
    def __init__(self, repo_path: str):
        """
        Initialize the UAT documentation generator.
        
        Args:
            repo_path (str): Path to the repository
        """
        self.repo_path = Path(repo_path)
        if not self.repo_path.exists():
            raise ValueError(f"Repository path does not exist: {repo_path}")
        
        # Create documentation directories if they don't exist
        self.logic_docs_dir = self.repo_path / 'Logic Understanding'
        if not self.logic_docs_dir.exists():
            raise ValueError("Logic Understanding directory not found")
            
        self.uat_docs_dir = self.repo_path / 'UAT Documentation'
        self.uat_docs_dir.mkdir(exist_ok=True)
        
        # Initialize OpenAI client
        self._setup_openai()
        
        # Initialize Git repository for versioning
        try:
            self.repo = git.Repo(self.repo_path)
        except git.InvalidGitRepositoryError:
            logger.warning("Not a git repository. Versioning will be limited.")
            self.repo = None
        
        # Store analysis results
        self.project_doc: Optional[Document] = None
        self.test_cases: Dict = {}
        self.test_environments: Dict = {}
        
    def _setup_openai(self):
        """Setup OpenAI API with proper error handling."""
        script_dir = Path(__file__).parent.absolute()
        env_path = script_dir / '.env'
        
        if not env_path.exists():
            raise FileNotFoundError(f".env file not found at: {env_path}")
        
        load_dotenv(env_path)
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            raise ValueError("OPENAI_API_KEY not found in .env file")
        
        if not api_key.startswith('sk-'):
            logger.warning("OPENAI_API_KEY doesn't appear to be in the correct format")
        
        self.client = openai.OpenAI(api_key=api_key)
        logger.info("Successfully initialized OpenAI client")

    def _read_project_documentation(self) -> bool:
        """Read the latest project logic documentation."""
        try:
            doc_path = self.logic_docs_dir / "project_logic_documentation_latest.docx"
            if not doc_path.exists():
                raise FileNotFoundError("Latest project documentation not found")
            
            self.project_doc = Document(doc_path)
            logger.info("Successfully loaded project documentation")
            return True
            
        except Exception as e:
            logger.error(f"Error reading project documentation: {str(e)}")
            return False

    def _extract_document_content(self) -> Dict:
        """Extract content from the project documentation."""
        if not self.project_doc:
            raise ValueError("Project documentation not loaded")
        
        content = {
            'project_description': '',
            'core_functionality': [],
            'key_features': []
        }
        
        current_section = None
        
        for paragraph in self.project_doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Check for section headers
            if paragraph.style.name.startswith('Heading 1'):
                if 'Project Description' in text:
                    current_section = 'project_description'
                elif 'Core Functionality' in text:
                    current_section = 'core_functionality'
                elif 'Key Features' in text:
                    current_section = 'key_features'
                continue
            
            # Add content to appropriate section
            if current_section == 'project_description':
                content['project_description'] += text + '\n'
            elif current_section == 'core_functionality' and text.startswith('•'):
                content['core_functionality'].append(text[1:].strip())
            elif current_section == 'key_features' and text.startswith('•'):
                content['key_features'].append(text[1:].strip())
        
        return content

    def _generate_test_cases(self, content: Dict) -> Dict:
        """Generate test cases using GPT-4."""
        try:
            prompt = f"""Based on this project documentation, generate comprehensive test cases and test environments.
Return a JSON object with these exact keys:
{{
    "test_environments": [
        {{
            "name": "string describing the environment",
            "description": "string describing the environment setup",
            "requirements": ["list", "of", "environment requirements"],
            "setup_steps": ["list", "of", "setup steps"]
        }}
    ],
    "test_cases": [
        {{
            "id": "unique test case ID",
            "title": "string describing the test case",
            "description": "string describing what is being tested",
            "environment": "string matching one of the test environment names",
            "priority": "High/Medium/Low",
            "category": "Functional/Integration/System/Performance/Security",
            "preconditions": ["list", "of", "preconditions"],
            "steps": ["list", "of", "test steps"],
            "expected_results": ["list", "of", "expected results"],
            "dependencies": ["list", "of", "test case IDs this depends on"]
        }}
    ]
}}

Project Documentation:
{json.dumps(content, indent=2)}

IMPORTANT: 
1. Return ONLY a valid JSON object with the exact keys shown above
2. Generate test cases for all core functionality and key features
3. Include at least one test environment
4. Test cases should be comprehensive and cover all major functionality
5. Each test case should be specific and verifiable"""

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert QA engineer. Generate comprehensive test cases and environments based on project documentation. Return ONLY valid JSON with the exact structure specified."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4000,
                response_format={ "type": "json_object" }
            )
            
            try:
                test_data = json.loads(response.choices[0].message.content)
                
                # Validate required keys
                required_keys = {'test_environments', 'test_cases'}
                if not all(key in test_data for key in required_keys):
                    raise ValueError("Missing required keys in response")
                
                # Validate test environments
                if not isinstance(test_data['test_environments'], list):
                    raise ValueError("test_environments must be a list")
                
                # Validate test cases
                if not isinstance(test_data['test_cases'], list):
                    raise ValueError("test_cases must be a list")
                
                # Validate each test case
                for tc in test_data['test_cases']:
                    required_tc_keys = {
                        'id', 'title', 'description', 'environment',
                        'priority', 'category', 'preconditions',
                        'steps', 'expected_results', 'dependencies'
                    }
                    if not all(key in tc for key in required_tc_keys):
                        raise ValueError(f"Test case {tc.get('id', 'unknown')} missing required keys")
                
                return test_data
                
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON in test cases response: {str(e)}")
                logger.error(f"Raw response: {response.choices[0].message.content}")
                return {
                    'test_environments': [],
                    'test_cases': []
                }
            
        except Exception as e:
            logger.error(f"Error generating test cases: {str(e)}")
            return {
                'test_environments': [],
                'test_cases': []
            }

    def _create_uat_document(self, content: Dict, test_data: Dict) -> Document:
        """Create the UAT documentation document."""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "User Acceptance Testing Documentation"
        doc.core_properties.author = "UAT Documentation Generator"
        
        # Add title
        title = doc.add_heading('User Acceptance Testing Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add version information
        if self.repo:
            version = doc.add_paragraph()
            version.add_run('Version: ').bold = True
            version.add_run(f"{self.repo.head.commit.hexsha[:8]}\n")
            version.add_run('Last Updated: ').bold = True
            version.add_run(f"{datetime.fromtimestamp(self.repo.head.commit.committed_date).strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add original project documentation
        doc.add_heading('Original Project Documentation', level=1)
        
        # Project Description
        doc.add_heading('Project Description', level=2)
        p = doc.add_paragraph()
        p.add_run(content['project_description'].strip())
        
        # Core Functionality
        doc.add_heading('Core Functionality', level=2)
        for func in content['core_functionality']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(func.strip())
        
        # Key Features
        doc.add_heading('Key Features', level=2)
        for feature in content['key_features']:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(feature.strip())
        
        # Add test environments
        doc.add_heading('Test Environments', level=1)
        for env in test_data['test_environments']:
            doc.add_heading(env['name'], level=2)
            p = doc.add_paragraph()
            p.add_run(env['description'].strip())
            
            # Requirements
            doc.add_heading('Requirements', level=3)
            for req in env['requirements']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(req.strip())
            
            # Setup Steps
            doc.add_heading('Setup Steps', level=3)
            for step in env['setup_steps']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(step.strip())
        
        # Add test cases
        doc.add_heading('Test Cases', level=1)
        
        # Group test cases by category
        categories = {}
        for tc in test_data['test_cases']:
            cat = tc['category']
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(tc)
        
        # Add test cases by category
        for category, test_cases in categories.items():
            doc.add_heading(f'{category} Test Cases', level=2)
            
            for tc in test_cases:
                # Test Case Header
                header = doc.add_heading(f"Test Case {tc['id']}: {tc['title']}", level=3)
                header.style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                
                # Test Case Details
                p = doc.add_paragraph()
                p.add_run('Description: ').bold = True
                p.add_run(tc['description'].strip())
                
                p = doc.add_paragraph()
                p.add_run('Environment: ').bold = True
                p.add_run(tc['environment'].strip())
                
                p = doc.add_paragraph()
                p.add_run('Priority: ').bold = True
                p.add_run(tc['priority'].strip())
                
                # Preconditions
                doc.add_heading('Preconditions', level=4)
                for pre in tc['preconditions']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(pre.strip())
                
                # Test Steps
                doc.add_heading('Test Steps', level=4)
                for i, step in enumerate(tc['steps'], 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(step.strip())
                
                # Expected Results
                doc.add_heading('Expected Results', level=4)
                for result in tc['expected_results']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(result.strip())
                
                # Dependencies
                if tc['dependencies']:
                    doc.add_heading('Dependencies', level=4)
                    for dep in tc['dependencies']:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"Test Case {dep}".strip())
                
                doc.add_paragraph()  # Add spacing between test cases
        
        return doc

    def generate_uat_documentation(self) -> str:
        """Generate and save the UAT documentation."""
        logger.info("Generating UAT documentation...")
        
        # Read project documentation
        if not self._read_project_documentation():
            raise ValueError("Failed to read project documentation")
        
        # Extract content
        content = self._extract_document_content()
        
        # Generate test cases
        test_data = self._generate_test_cases(content)
        
        # Create UAT document
        doc = self._create_uat_document(content, test_data)
        
        # Generate filename with version
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if self.repo:
            version = self.repo.head.commit.hexsha[:8]
            filename = f"uat_documentation_v1_{version}_{timestamp}.docx"
        else:
            filename = f"uat_documentation_v1_{timestamp}.docx"
        
        # Save document in UAT Documentation folder
        doc_path = self.uat_docs_dir / filename
        doc.save(doc_path)
        
        # Create latest symlink/copy
        latest_doc = self.uat_docs_dir / "uat_documentation_latest.docx"
        if os.name == 'nt':
            shutil.copy2(doc_path, latest_doc)
        else:
            if latest_doc.exists():
                latest_doc.unlink()
            latest_doc.symlink_to(doc_path)
        
        logger.info(f"UAT documentation saved to {doc_path}")
        return filename

def main():
    """Main function to generate UAT documentation."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate UAT documentation from project logic documentation')
    parser.add_argument('repo_path', help='Path to the repository')
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose logging')
    
    args = parser.parse_args()
    
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    
    try:
        generator = UATDocumentationGenerator(args.repo_path)
        doc_file = generator.generate_uat_documentation()
        
        print("\nUAT Documentation Generated Successfully!")
        print("=" * 50)
        print(f"Documentation file: UAT Documentation/{doc_file}")
        print(f"Latest version: UAT Documentation/uat_documentation_latest.docx")
        print("\nThe documentation includes:")
        print("• Original Project Documentation")
        print("• Test Environments")
        print("• Test Cases (by category)")
        
    except Exception as e:
        logger.error(f"Error generating UAT documentation: {str(e)}")
        raise

if __name__ == '__main__':
    main() 