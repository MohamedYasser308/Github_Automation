#!/usr/bin/env python3

import os
import json
import logging
from pathlib import Path
from typing import Dict, List, Set, Optional, Tuple
import re
from collections import defaultdict
from datetime import datetime
import glob
import mimetypes  # Use built-in mimetypes instead of python-magic
import chardet  # for encoding detection
import hashlib
from dataclasses import dataclass, asdict
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Initialize mimetypes
mimetypes.init()

@dataclass
class FileClassification:
    """Data class to store file classification information."""
    filename: str
    relative_path: str
    file_type: str
    mime_type: str
    encoding: str
    size: int
    md5_hash: str
    language: Optional[str] = None
    framework: Optional[str] = None
    dependencies: List[str] = None
    imports: List[str] = None
    complexity: Optional[int] = None
    last_modified: str = None
    created: str = None
    lines_of_code: Optional[int] = None
    comment_lines: Optional[int] = None
    blank_lines: Optional[int] = None

class FileClassifier:
    """Classify individual files in a repository."""
    
    # Language detection patterns
    LANGUAGE_PATTERNS = {
        'Python': {
            'extensions': {'.py', '.pyw', '.pyi'},
            'patterns': [r'^#!.*python', r'import\s+\w+', r'from\s+\w+\s+import'],
            'frameworks': {
                'django': [r'from\s+django\.', r'Django\s+settings'],
                'flask': [r'from\s+flask\s+import', r'Flask\('],
                'fastapi': [r'from\s+fastapi\s+import', r'FastAPI\('],
            }
        },
        'JavaScript': {
            'extensions': {'.js', '.jsx', '.mjs'},
            'patterns': [r'^//', r'^/\*', r'function\s+\w+', r'const\s+\w+', r'let\s+\w+'],
            'frameworks': {
                'react': [r'import\s+React', r'from\s+\'react\'', r'ReactDOM\.render'],
                'vue': [r'import\s+Vue', r'new\s+Vue\('],
                'angular': [r'@angular', r'NgModule'],
            }
        },
        'TypeScript': {
            'extensions': {'.ts', '.tsx'},
            'patterns': [r'interface\s+\w+', r'type\s+\w+', r'import\s+type'],
            'frameworks': {
                'react': [r'import\s+React', r'from\s+\'react\''],
                'angular': [r'@angular', r'NgModule'],
            }
        },
        'Java': {
            'extensions': {'.java'},
            'patterns': [r'public\s+class', r'import\s+java\.', r'@Override'],
            'frameworks': {
                'spring': [r'@SpringBootApplication', r'@Controller', r'@Service'],
                'hibernate': [r'@Entity', r'@Table'],
            }
        },
        'Ruby': {
            'extensions': {'.rb'},
            'patterns': [r'require\s+\'', r'class\s+\w+', r'def\s+\w+'],
            'frameworks': {
                'rails': [r'class\s+\w+Controller', r'ActiveRecord::Base'],
            }
        },
        'PHP': {
            'extensions': {'.php'},
            'patterns': [r'<\?php', r'namespace\s+\w+', r'use\s+\w+'],
            'frameworks': {
                'laravel': [r'use\s+Illuminate\\', r'Route::'],
                'symfony': [r'use\s+Symfony\\', r'@Route'],
            }
        },
        'Go': {
            'extensions': {'.go'},
            'patterns': [r'package\s+\w+', r'import\s+\('],
            'frameworks': {
                'gin': [r'gin\.', r'github\.com/gin-gonic/gin'],
            }
        },
    }

    def __init__(self, repo_path: str):
        """
        Initialize the file classifier.
        
        Args:
            repo_path (str): Path to the repository to analyze
        """
        self.repo_path = Path(repo_path)
        if not self.repo_path.exists():
            raise ValueError(f"Repository path does not exist: {repo_path}")
        
        self.classifier_dir = self.repo_path / 'Classifier'
        self.classifier_dir.mkdir(exist_ok=True)

    def _get_version_info(self) -> Tuple[int, str]:
        """Get the next version number and timestamp for classification files."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        existing_files = glob.glob(str(self.classifier_dir / "file_classification_v*.json"))
        if not existing_files:
            return 1, timestamp
            
        versions = []
        for file in existing_files:
            match = re.search(r'file_classification_v(\d+)_', file)
            if match:
                versions.append(int(match.group(1)))
        
        next_version = max(versions) + 1 if versions else 1
        return next_version, timestamp

    def _detect_language_and_framework(self, file_path: Path, content: str) -> Tuple[Optional[str], Optional[str], List[str]]:
        """Detect programming language and framework from file content."""
        ext = file_path.suffix.lower()
        imports = []
        framework = None
        
        # First check by extension
        for lang, info in self.LANGUAGE_PATTERNS.items():
            if ext in info['extensions']:
                # Check for framework patterns
                for fw, patterns in info['frameworks'].items():
                    if any(re.search(pattern, content) for pattern in patterns):
                        framework = fw
                        break
                
                # Extract imports
                if lang == 'Python':
                    imports = re.findall(r'import\s+(\w+)', content) + re.findall(r'from\s+(\w+)\s+import', content)
                elif lang in ['JavaScript', 'TypeScript']:
                    imports = re.findall(r'import\s+.*?from\s+[\'"]([^\'"]+)[\'"]', content)
                elif lang == 'Java':
                    imports = re.findall(r'import\s+([\w.]+)', content)
                
                return lang, framework, imports
                
        return None, None, []

    def _analyze_file_content(self, file_path: Path) -> Tuple[int, int, int]:
        """Analyze file content for lines of code, comments, and blank lines."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.readlines()
            
            total_lines = len(content)
            blank_lines = sum(1 for line in content if not line.strip())
            comment_lines = 0
            
            ext = file_path.suffix.lower()
            if ext in {'.py', '.pyw', '.pyi'}:
                comment_lines = sum(1 for line in content if line.strip().startswith('#'))
            elif ext in {'.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h'}:
                comment_lines = sum(1 for line in content if line.strip().startswith('//'))
                # Count multi-line comments
                in_comment = False
                for line in content:
                    if '/*' in line:
                        in_comment = True
                    if '*/' in line:
                        in_comment = False
                        comment_lines += 1
                    elif in_comment:
                        comment_lines += 1
            
            return total_lines, comment_lines, blank_lines
            
        except Exception as e:
            logger.error(f"Error analyzing file content for {file_path}: {str(e)}")
            return 0, 0, 0

    def _calculate_complexity(self, file_path: Path, content: str) -> int:
        """Calculate cyclomatic complexity of the file."""
        complexity = 1  # Base complexity
        
        # Common control flow statements that increase complexity
        patterns = [
            r'\bif\b', r'\belse\b', r'\bfor\b', r'\bwhile\b',
            r'\bcase\b', r'\bcatch\b', r'\b&&\b', r'\b\|\|\b',
            r'\?', r':', r'\breturn\b'
        ]
        
        for pattern in patterns:
            complexity += len(re.findall(pattern, content))
            
        return complexity

    def classify_file(self, file_path: Path) -> FileClassification:
        """Classify a single file."""
        try:
            # Get file metadata
            stat = file_path.stat()
            
            # Get MIME type using mimetypes
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type is None:
                # Fallback to extension-based detection
                ext = file_path.suffix.lower()
                if ext in {'.py', '.pyw', '.pyi'}:
                    mime_type = 'text/x-python'
                elif ext in {'.js', '.jsx', '.mjs'}:
                    mime_type = 'application/javascript'
                elif ext in {'.ts', '.tsx'}:
                    mime_type = 'application/typescript'
                elif ext in {'.java'}:
                    mime_type = 'text/x-java-source'
                elif ext in {'.rb'}:
                    mime_type = 'text/x-ruby'
                elif ext in {'.php'}:
                    mime_type = 'application/x-httpd-php'
                elif ext in {'.go'}:
                    mime_type = 'text/x-go'
                elif ext in {'.html', '.htm'}:
                    mime_type = 'text/html'
                elif ext in {'.css'}:
                    mime_type = 'text/css'
                elif ext in {'.json'}:
                    mime_type = 'application/json'
                elif ext in {'.xml'}:
                    mime_type = 'application/xml'
                elif ext in {'.md', '.markdown'}:
                    mime_type = 'text/markdown'
                elif ext in {'.txt'}:
                    mime_type = 'text/plain'
                else:
                    mime_type = 'application/octet-stream'
            
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                encoding = result['encoding'] or 'unknown'
            
            # Calculate MD5 hash
            md5_hash = hashlib.md5(raw_data).hexdigest()
            
            # Read file content for analysis
            try:
                content = raw_data.decode(encoding, errors='ignore')
            except:
                content = raw_data.decode('utf-8', errors='ignore')
            
            # Detect language and framework
            language, framework, imports = self._detect_language_and_framework(file_path, content)
            
            # Analyze content
            total_lines, comment_lines, blank_lines = self._analyze_file_content(file_path)
            complexity = self._calculate_complexity(file_path, content) if language else None
            
            return FileClassification(
                filename=file_path.name,
                relative_path=str(file_path.relative_to(self.repo_path)),
                file_type=file_path.suffix.lower(),
                mime_type=mime_type,
                encoding=encoding,
                size=stat.st_size,
                md5_hash=md5_hash,
                language=language,
                framework=framework,
                dependencies=imports,
                imports=imports,
                complexity=complexity,
                last_modified=datetime.fromtimestamp(stat.st_mtime).isoformat(),
                created=datetime.fromtimestamp(stat.st_ctime).isoformat(),
                lines_of_code=total_lines,
                comment_lines=comment_lines,
                blank_lines=blank_lines
            )
            
        except Exception as e:
            logger.error(f"Error classifying file {file_path}: {str(e)}")
            return None

    def _create_word_document(self, classification_data: Dict) -> Document:
        """
        Create a Word document with the classification results.
        
        Args:
            classification_data (Dict): Complete classification data
            
        Returns:
            Document: Word document object
        """
        doc = Document()
        
        # Set document properties
        repo_name = Path(classification_data['repository']).name
        doc.core_properties.title = f"File Classification Report: {repo_name} (v{classification_data['version']})"
        doc.core_properties.author = "File Classifier"
        doc.core_properties.created = datetime.now()
        
        # Add title
        title = doc.add_heading('File Classification Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add repository info and version
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f'Repository: {repo_name}\n').bold = True
        subtitle.add_run(f'Version: {classification_data["version"]}\n').bold = True
        subtitle.add_run(f'Analysis Date: {datetime.strptime(classification_data["timestamp"], "%Y%m%d_%H%M%S").strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Add summary section
        doc.add_heading('Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run(f'Total Files Classified: {classification_data["total_files"]}\n\n').bold = True
        
        # Count and display languages
        languages = defaultdict(int)
        frameworks = defaultdict(int)
        file_types = defaultdict(int)
        total_size = 0
        total_lines = 0
        
        for file in classification_data['classifications']:
            if file['language']:
                languages[file['language']] += 1
            if file['framework']:
                frameworks[file['framework']] += 1
            file_types[file['file_type']] += 1
            total_size += file['size']
            if file['lines_of_code']:
                total_lines += file['lines_of_code']
        
        # Add statistics
        doc.add_heading('Repository Statistics', level=2)
        stats = doc.add_paragraph()
        stats.add_run(f'Total Repository Size: {total_size / 1024 / 1024:.2f} MB\n')
        stats.add_run(f'Total Lines of Code: {total_lines:,}\n')
        
        # Add languages section
        doc.add_heading('Languages', level=2)
        if languages:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Language'
            header_cells[1].text = 'Files'
            
            for lang, count in sorted(languages.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = lang
                row_cells[1].text = str(count)
        else:
            doc.add_paragraph('No programming languages detected')
        
        # Add frameworks section
        doc.add_heading('Frameworks', level=2)
        if frameworks:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Framework'
            header_cells[1].text = 'Files'
            
            for fw, count in sorted(frameworks.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = fw
                row_cells[1].text = str(count)
        else:
            doc.add_paragraph('No frameworks detected')
        
        # Add file types section
        doc.add_heading('File Types', level=2)
        if file_types:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'File Type'
            header_cells[1].text = 'Count'
            
            for ext, count in sorted(file_types.items()):
                row_cells = table.add_row().cells
                row_cells[0].text = ext or '(no extension)'
                row_cells[1].text = str(count)
        
        # Add detailed file analysis
        doc.add_heading('Detailed File Analysis', level=1)
        
        # Group files by language
        files_by_language = defaultdict(list)
        for file in classification_data['classifications']:
            if file['language']:
                files_by_language[file['language']].append(file)
            else:
                files_by_language['Other'].append(file)
        
        for language, files in sorted(files_by_language.items()):
            doc.add_heading(f'{language} Files', level=2)
            
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'File'
            header_cells[1].text = 'Framework'
            header_cells[2].text = 'Size'
            header_cells[3].text = 'Lines'
            header_cells[4].text = 'Complexity'
            header_cells[5].text = 'Last Modified'
            
            for file in sorted(files, key=lambda x: x['relative_path']):
                row_cells = table.add_row().cells
                row_cells[0].text = file['relative_path']
                row_cells[1].text = file['framework'] or '-'
                row_cells[2].text = f"{file['size'] / 1024:.1f} KB"
                row_cells[3].text = str(file['lines_of_code'] or '-')
                row_cells[4].text = str(file['complexity'] or '-')
                row_cells[5].text = datetime.fromisoformat(file['last_modified']).strftime('%Y-%m-%d')
        
        # Add footer with timestamp
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc

    def classify_repository(self) -> Dict:
        """Classify all files in the repository."""
        logger.info(f"Classifying files in repository: {self.repo_path}")
        
        classifications = []
        skipped_dirs = {'.git', 'node_modules', '__pycache__', 'venv', '.idea', 'Classifier'}
        
        for root, dirs, files in os.walk(self.repo_path):
            # Skip certain directories
            if any(skip in root for skip in skipped_dirs):
                continue
                
            for file in files:
                file_path = Path(root) / file
                
                # Skip hidden files and files in the Classifier directory
                if file.startswith('.') or 'Classifier' in str(file_path):
                    continue
                
                classification = self.classify_file(file_path)
                if classification:
                    classifications.append(asdict(classification))
        
        # Get version info
        version, timestamp = self._get_version_info()
        
        # Prepare the complete classification data
        classification_data = {
            'version': version,
            'timestamp': timestamp,
            'repository': str(self.repo_path),
            'total_files': len(classifications),
            'classifications': classifications
        }
        
        # Save versioned JSON file
        json_filename = f"file_classification_v{version}_{timestamp}.json"
        json_path = self.classifier_dir / json_filename
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(classification_data, f, indent=2)
        
        # Generate and save Word document
        doc = self._create_word_document(classification_data)
        docx_filename = f"file_classification_v{version}_{timestamp}.docx"
        docx_path = self.classifier_dir / docx_filename
        doc.save(str(docx_path))
        
        # Create latest version links/copies
        latest_json = self.classifier_dir / "file_classification_latest.json"
        latest_docx = self.classifier_dir / "file_classification_latest.docx"
        
        if os.name == 'nt':
            shutil.copy2(json_path, latest_json)
            shutil.copy2(docx_path, latest_docx)
        else:
            if latest_json.exists():
                latest_json.unlink()
            if latest_docx.exists():
                latest_docx.unlink()
            latest_json.symlink_to(json_path)
            latest_docx.symlink_to(docx_path)
        
        logger.info(f"Classification results saved to {self.classifier_dir}")
        logger.info(f"Version {version} created at {timestamp}")
        
        return classification_data

def main():
    """Main function to classify repository files."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Classify files in a repository')
    parser.add_argument('repo_path', help='Path to the repository to analyze')
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose logging')
    parser.add_argument('--list-versions', '-l', action='store_true', help='List existing classification versions')
    
    args = parser.parse_args()
    
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    
    try:
        classifier = FileClassifier(args.repo_path)
        
        # List versions if requested
        if args.list_versions:
            classifier_dir = Path(args.repo_path) / 'Classifier'
            if classifier_dir.exists():
                files = glob.glob(str(classifier_dir / "file_classification_v*.json"))
                if files:
                    print("\nExisting Classification Versions:")
                    print("=" * 50)
                    for file in sorted(files):
                        match = re.search(r'file_classification_v(\d+)_(\d{8}_\d{6})', file)
                        if match:
                            version, timestamp = match.groups()
                            dt = datetime.strptime(timestamp, "%Y%m%d_%H%M%S")
                            print(f"Version {version}: {dt.strftime('%Y-%m-%d %H:%M:%S')} ({Path(file).name})")
                else:
                    print("No previous classification versions found.")
            else:
                print("No Classifier directory found.")
            return
        
        # Perform classification
        result = classifier.classify_repository()
        
        # Print summary
        print("\nRepository Classification Summary:")
        print("=" * 50)
        print(f"Total Files Classified: {result['total_files']}")
        
        # Count languages
        languages = defaultdict(int)
        frameworks = defaultdict(int)
        for file in result['classifications']:
            if file['language']:
                languages[file['language']] += 1
            if file['framework']:
                frameworks[file['framework']] += 1
        
        print("\nLanguages Found:")
        for lang, count in sorted(languages.items()):
            print(f"- {lang}: {count} files")
            
        print("\nFrameworks Found:")
        for fw, count in sorted(frameworks.items()):
            print(f"- {fw}: {count} files")
        
        print(f"\nClassification results have been saved in the repository's 'Classifier' directory:")
        print(f"- Latest JSON report: Classifier/file_classification_latest.json")
        print(f"- Latest Word document: Classifier/file_classification_latest.docx")
        print("\nUse --list-versions to see all classification versions")
            
    except Exception as e:
        logger.error(f"Error classifying repository: {str(e)}")
        raise

if __name__ == '__main__':
    main() 